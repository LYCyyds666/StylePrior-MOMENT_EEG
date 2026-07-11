"""
Lightweight consistency audit for the StylePrior-MOMENT manuscript.

Usage:
    python audit_paper_consistency.py StylePrior-MOMENT_EEG_paper_revised_references.docx
"""
from __future__ import annotations
import re, sys
from pathlib import Path
from docx import Document

BAD_PATTERNS = {
    'old_method_name': [r'\bSS-MOMENT\b', r'\bSAGESTREAM\b', r'\bSageStream\b'],
    'wrong_stsa_story': [r'confidence-guided', r'confidence score', r'reliable samples', r'down-?weighted', r'agreement between incoming', r'more consistent with the current style estimate'],
    'overclaim': [r'consistently achieves robust', r'ensuring robust generalization', r'each component contributes', r'reliable gain'],
    'shared_module_misstatement': [r'same SA-MoE module is shared', r'SA-MoE module is shared across all'],
    'canonical_style_risk': [r'same canonical style', r'canonical style for every subject'],
}

REQUIRED_PHRASES = [
    'StylePrior-MOMENT',
    'batch-wise streaming adaptation',
    'discrepancy weight',
    'source-derived style',
    'target-adaptive',
]


def get_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(parts)


def main():
    if len(sys.argv) != 2:
        print('Usage: python audit_paper_consistency.py manuscript.docx')
        sys.exit(2)
    path = Path(sys.argv[1])
    text = get_text(path)
    print(f'Auditing: {path}')
    print('='*70)
    total_hits = 0
    for group, patterns in BAD_PATTERNS.items():
        print(f'[{group}]')
        group_hits = 0
        for pat in patterns:
            hits = list(re.finditer(pat, text, flags=re.I))
            if hits:
                group_hits += len(hits)
                total_hits += len(hits)
                print(f'  {pat}: {len(hits)} hit(s)')
                for h in hits[:3]:
                    s=max(0,h.start()-90); e=min(len(text),h.end()+90)
                    print('   ...' + text[s:e].replace('\n',' ') + '...')
        if group_hits == 0:
            print('  OK')
    print('\n[required phrases]')
    for ph in REQUIRED_PHRASES:
        print(f'  {ph}:', 'OK' if ph.lower() in text.lower() else 'MISSING')
    print('\nSummary:', 'PASS' if total_hits == 0 else f'{total_hits} possible issue(s) found')

if __name__ == '__main__':
    main()
